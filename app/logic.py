import re
from datetime import datetime

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from app.config import Config


def read_excel(file_path):
    """读取Excel文件并返回工作簿对象"""
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    return workbook


def extract_subjects(workbook):
    """从目录sheet页提取科目名称"""
    sheet = workbook['目录']
    subjects = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[3] == '是':  # 假设"是否做附注表格"在第2列
            subjects.append(row[1])  # 假设"科目名称"在第1列
    return subjects


def format_number(value):
    """格式化数字，添加千位分隔符"""
    if isinstance(value, (int, float)):
        return '{:,.2f}'.format(value)
    return value


def extract_table_data(workbook, sheet_name):
    """提取财务表格数据并进行格式化"""
    sheet = workbook[sheet_name]
    data = []
    for row in sheet.iter_rows(min_row=3, min_col=3, values_only=True):  # 从第3行第3列开始
        formatted_row = [format_number(cell) if cell is not None else "" for cell in row]
        data.append(formatted_row)
    return data


def set_cell_font(cell, font_name, font_size, bold=False):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.bold = bold


def set_table_borders(table):
    """设置表格的上下两条外边框加粗，左右两条外边框不显示"""
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    # 设置上下边框加粗
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    tblBorders.append(top)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    tblBorders.append(bottom)

    # 设置左右边框不显示
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'none')
    tblBorders.append(left)

    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'none')
    tblBorders.append(right)

    # 其他边框设置为默认
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    tblBorders.append(insideH)

    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'single')
    tblBorders.append(insideV)

    tbl.tblPr.append(tblBorders)


def insert_table_into_doc(doc, table_data, replacements, sheet_name, paragraph):
    """将表格插入到Word文档的指定位置"""
    balance = calculate_balance(table_data)
    report_date = replacements.get('报表截止日', '')
    sentence = f"截止{report_date}，公司{sheet_name}账面余额为{balance}元。"

    # 创建表格段落
    tbl = doc.add_table(rows=0, cols=len(table_data[0]))
    tbl.style = 'Table Grid'
    tbl.autofit = False

    # 插入表格数据
    for row_idx, row_data in enumerate(table_data):
        row = tbl.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = str(cell_data)

            # 表头样式设置
            if row_idx == 0:
                set_cell_font(row[i], '宋体', 10, bold=True)
                row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
            # 表格最后一行如果包含“合计”，也要加粗且居中显示
            elif cell_data and "合计" in str(cell_data):
                set_cell_font(row[i], '宋体', 10, bold=True)
                row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
            else:
                if isinstance(cell_data, (int, float)):
                    set_cell_font(row[i], 'Arial Narrow', 10)
                else:
                    set_cell_font(row[i], '宋体', 10)

    # 设置表格的上下两条外边框加粗，左右两条外边框不显示
    set_table_borders(tbl)

    # 插入表格段落到句子段落下方
    paragraph._element.addnext(tbl._element)

    # 创建句子段落并插入到表格前面
    sentence_paragraph = doc.add_paragraph(sentence)
    sentence_paragraph.paragraph_format.first_line_indent = Inches(0.2)
    sentence_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    sentence_paragraph.paragraph_format.line_spacing = Pt(22)
    sentence_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    tbl._element.addprevious(sentence_paragraph._element)
def calculate_balance(table_data):
    """计算表格账面余额，如果最后一行有合计并且表格最后一列是期末余额，则取合计行的最右侧单元格，否则为0.00"""
    balance = 0.0

    # 检查最后一行是否包含“合计”并且最后一列是“期末余额”
    if len(table_data) > 1 and "合计" in str(table_data[-1][0]) and "期末金额" in str(table_data[0][-1]):
        try:
            balance_str = table_data[-1][-1].replace(',', '')  # 去掉逗号
            balance = float(balance_str)
        except (ValueError, TypeError):
            balance = 0.0
    else:
        balance = 0.0

    return f'{balance:,.2f}'

def clean_placeholder(text):
    """去除占位符中的空格"""
    return re.sub(r'\s+', '', text)


def extract_key_value_data(file_path, sheet_name):
    """读取指定Excel文件和sheet页中的数据，返回<key, value>数据结构"""
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    sheet = workbook[sheet_name]
    data = {}

    for row in sheet.iter_rows(min_row=2, min_col=3, values_only=False):
        key_cell = row[0]
        value_cell = row[1]

        # 处理key
        key = key_cell.value
        if key is not None:
            key = str(key).strip()
            if key == '审计报告编号':
                key = '报告编号'

        # 处理value
        value = value_cell.value
        if value is not None:
            if isinstance(value, datetime):
                value = value.strftime('%Y-%m-%d')
                if key == '报表截止日':
                    data['报告年度'] = value[:4] + '年'  # 提取年份
            elif isinstance(value, float):
                value = f'{value:.2f}'
            else:
                value = str(value).strip()

        if key and value:
            data[key] = value

    # 特殊处理报告日
    if '审计报告日期' in data:
        data['报告日'] = convert_date_to_chinese(data['审计报告日期'])

    # 拼接企业信息为一个字符串
    company_info_parts = [
        data.get('企业全称', '') + "（以下简称“本公司”或“公司”），于" + data.get('企业成立日期',
                                                                              '') + "成立，社会统一信用代码为" + data.get(
            '营业执照号码', '') + "，发证机关为" + data.get('批准工商行政机关', '') + "。",
        "法定代表人：" + data.get('法定代表人', ''),
        "注册资本：人民币" + data.get('注册资本', '') + "万元",
        "住所地：" + data.get('住所地', ''),
        "经营范围：" + data.get('经营范围', '') + "。"
    ]

    # 将多个部分合并为一个字符串，使用特定标记进行分隔
    data['企业信息'] = "##".join(company_info_parts)

    return data


def replace_placeholder_in_paragraph(paragraph, replacements):
    """在段落中用replacements中的值替换«key»结构的占位符，并保持原有的样式"""
    for key, value in replacements.items():
        # placeholder = f'«{key}»'
        cleaned_placeholder = clean_placeholder(key)

        runs = paragraph.runs
        for i in range(1, len(runs) - 1):
            prev_run_text = clean_placeholder(runs[i - 1].text.strip())
            current_run_text = clean_placeholder(runs[i].text.strip())
            next_run_text = clean_placeholder(runs[i + 1].text.strip())

            if prev_run_text == '«' and current_run_text == cleaned_placeholder and next_run_text == '»':
                runs[i - 1].text = ''
                if cleaned_placeholder == '企业信息':
                    # 插入新段落
                    parts = value.split('##')
                    for part in parts:
                        new_paragraph = paragraph.insert_paragraph_before(part)
                        new_paragraph.paragraph_format.first_line_indent = Pt(22)
                        new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        new_paragraph.paragraph_format.line_spacing = Pt(22)
                    runs[i].text = ''
                else:
                    runs[i].text = value
                runs[i + 1].text = ''


def replace_placeholders_in_doc(doc, replacements, sheet_names, workbook):
    """遍历Word文档的所有部分，替换占位符，并清空指定关键词之间的段落，插入sheet页名称和对应的财务表格"""
    # 清空“五、财务报表主要项目注释”和“六、或有事项”之间的所有段落
    clear_paragraphs_between_keywords(doc, "五、财务报表主要项目注释", "六、或有事项")

    # 插入sheet页名称和对应的财务表格
    insert_sheet_names_and_tables(doc, "五、财务报表主要项目注释", "六、或有事项", sheet_names, workbook, replacements)

    # 替换占位符
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_placeholder_in_paragraph(paragraph, replacements)

        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_placeholder_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, replacements)


def export_report(template, report_type, export_format):
    """导出报告"""
    excel_path = Config.EXCEL_FILE_PATH
    word_path = Config.TEMPLATE_PATHS.get((template, report_type))
    if not word_path:
        raise ValueError("未找到对应的报告模板")

    # 读取Excel文件
    workbook = read_excel(excel_path)

    # 提取科目名称
    subjects = extract_subjects(workbook)

    # 读取基本信息表数据
    key_value_data = extract_key_value_data(Config.BASIC_EXCEL_FILE_PATH, '基本信息表')

    # 读取Word文档
    doc = Document(word_path)

    # 替换Word文档中的占位符
    replace_placeholders_in_doc(doc, key_value_data,workbook.sheetnames,workbook)

    # 提取并插入表格数据
    # for subject in subjects:
    #     table_data = extract_table_data(workbook, subject)
    #     insert_table_into_doc(doc, table_data, f'«{subject}»',key_value_data)

    # 保存修改后的Word文档
    output_path = Config.OUT_WORD_FILE_PATH
    doc.save(output_path)

    if export_format == "PDF":
        # 如果需要将Word转换为PDF，可以使用第三方库如`python-docx2pdf`
        pass

    return output_path


def convert_date_to_chinese(date_str):
    """将日期字符串转换为中文大写日期格式"""
    chinese_numerals = {
        '0': '零', '1': '一', '2': '二', '3': '三', '4': '四', '5': '五',
        '6': '六', '7': '七', '8': '八', '9': '九'
    }
    year, month, day = date_str.split('-')

    def convert_year(y):
        return ''.join(chinese_numerals[digit] for digit in y) + '年'

    def convert_month(m):
        if m.startswith('0'):
            return chinese_numerals[m[1]] + '月'
        elif m == '10':
            return '十月'
        else:
            return '十' + chinese_numerals[m[1]] + '月' if m.startswith('1') else chinese_numerals[m[0]] + '月'

    def convert_day(d):
        if d.startswith('0'):
            return chinese_numerals[d[1]] + '日'
        elif d == '10':
            return '十日'
        elif d.startswith('1'):
            return '十' + chinese_numerals[d[1]] + '日'
        else:
            return '二十' + chinese_numerals[d[1]] + '日' if d.startswith('2') else '三十' + chinese_numerals[
                d[1]] + '日'

    return convert_year(year) + convert_month(month) + convert_day(day)

def clear_paragraphs_between_keywords(doc, start_keyword, end_keyword):
    """清空Word文档中两个关键词之间的所有段落"""
    start_found = False
    paragraphs_to_clear = []

    for paragraph in doc.paragraphs:
        if start_keyword in paragraph.text:
            start_found = True

        if start_found:
            paragraphs_to_clear.append(paragraph)

        if end_keyword in paragraph.text and start_found:
            break

    for paragraph in paragraphs_to_clear:
        if start_keyword not in paragraph.text and end_keyword not in paragraph.text:
            p = paragraph._element
            p.getparent().remove(p)
def add_paragraph_before(paragraph, text, indent=False):
    """在段落前添加一个新的段落，可以选择首行缩进"""
    new_paragraph = paragraph.insert_paragraph_before(text)
    if indent:
        new_paragraph.paragraph_format.first_line_indent = Inches(0.2)
    return new_paragraph


def insert_sheet_names_between_keywords(doc, start_keyword, end_keyword, sheet_names):
    """在两个关键词之间插入sheet页名称，并添加序号"""
    start_found = False
    paragraphs_to_insert = []

    for paragraph in doc.paragraphs:
        if start_keyword in paragraph.text:
            start_found = True

        if end_keyword in paragraph.text and start_found:
            for i, sheet_name in enumerate(sheet_names, start=1):
                new_paragraph = paragraph.insert_paragraph_before(f"{i}. {sheet_name}")
                new_paragraph.paragraph_format.first_line_indent = Inches(0.2)
                new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                new_paragraph.paragraph_format.line_spacing = Pt(22)
            break

def insert_sheet_names_and_tables(doc, start_keyword, end_keyword, sheet_names, workbook, replacements):
    """在两个关键词之间插入sheet页名称，并添加序号，同时插入对应的财务表格"""
    excluded_sheets = {"目录", "资表模板", "利表模板"}
    start_found = False
    counter = 1

    for paragraph in doc.paragraphs:
        if start_keyword in paragraph.text:
            start_found = True

        if end_keyword in paragraph.text and start_found:
            for sheet_name in sheet_names:
                if sheet_name not in excluded_sheets:
                    # 插入带序号的 sheet 页名称
                    new_paragraph = paragraph.insert_paragraph_before(f"{counter}、{sheet_name}")
                    new_paragraph.paragraph_format.first_line_indent = Inches(0.2)
                    new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    new_paragraph.paragraph_format.line_spacing = Pt(22)

                    # 插入模板替换后的句子和对应的财务表格
                    table_data = extract_table_data(workbook, sheet_name)
                    insert_table_into_doc(doc, table_data, replacements, sheet_name, new_paragraph)

                    counter += 1
            break